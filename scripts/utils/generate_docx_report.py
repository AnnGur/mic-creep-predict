"""
generate_docx_report.py — Generate KSE written report (.docx) for MIC Creep project.

Output: reports/MIC_Creep_Written_Report.docx

Run:
    .venv/bin/python scripts/utils/generate_docx_report.py
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import docx.opc.constants


PROJECT_ROOT = Path(__file__).resolve().parent.parent.parent
REPORTS = PROJECT_ROOT / "reports"
REPORTS.mkdir(parents=True, exist_ok=True)
OUT = REPORTS / "MIC_Creep_Report.docx"


def set_font(run, name="Calibri", size=11, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    set_font(run, size=13 if level == 1 else 11, bold=True,
             color=(31, 73, 125) if level == 1 else (0, 0, 0))
    return p


def body(doc, text, space_after=4):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(space_after)
    for run in p.runs:
        set_font(run)
    return p


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(10)
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if col_widths and i < len(col_widths):
            hdr_cells[i].width = Inches(col_widths[i])

    for row_data in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = str(val)
            row_cells[i].paragraphs[0].runs[0].font.size = Pt(10)
            row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if col_widths and i < len(col_widths):
                row_cells[i].width = Inches(col_widths[i])

    doc.add_paragraph()
    return table


def build_report():
    doc = Document()

    # ── Page margins ──────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # ── Title block ───────────────────────────────────────────────────────────
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run(
        "MIC Creep Prediction Using Machine Learning\n"
        "Vivli AMR Surveillance Data Challenge 2026"
    )
    set_font(title_run, size=16, bold=True, color=(31, 73, 125))
    title_p.paragraph_format.space_after = Pt(4)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_p.add_run(
        "Anna Gurina (Technical Lead) | Inna Kucherova (Research Lead / Domain Expert)\n"
        "KSE Introduction to Machine Learning - Group Project"
    )
    set_font(sub_run, size=11)
    sub_p.paragraph_format.space_after = Pt(8)

    # ── 1. Introduction ───────────────────────────────────────────────────────
    heading(doc, "1. Introduction")
    body(doc,
         "Antibiotic resistance represents one of the most urgent global health threats. "
         "The gradual upward drift of Minimum Inhibitory Concentration (MIC) values - "
         "termed 'MIC Creep' - is an early warning signal that precedes full clinical "
         "resistance. Because MIC values remain technically 'susceptible' while drifting "
         "upward, this phenomenon escapes detection by standard clinical breakpoint reporting. "
         "This project applies machine learning to quantify and predict MIC Creep for "
         "Klebsiella pneumoniae and Acinetobacter baumannii against Meropenem, a "
         "last-resort carbapenem antibiotic, using isolate-level surveillance data from "
         "the Vivli AMR Register (ATLAS dataset, Pfizer).")

    # ── 2. Problem Statement ──────────────────────────────────────────────────
    heading(doc, "2. Problem Statement")
    body(doc,
         "We frame MIC Creep as a regression problem: predict the log2-transformed MIC "
         "value for each bacterial isolate given its year, geographic origin, host "
         "demographics, and specimen source. The primary clinical question is: "
         "are MIC values systematically increasing over time, even within the susceptible "
         "range? We evaluate two organisms independently because they occupy distinct "
         "ecological niches: K. pneumoniae is predominantly a community/hospital-acquired "
         "pathogen, while A. baumannii is almost exclusively nosocomial and frequently "
         "carbapenem-resistant at baseline (pan-resistant outbreaks in ICUs).")

    heading(doc, "Why It Matters", level=2)
    body(doc,
         "Meropenem is a carbapenem antibiotic used as a last resort when all other "
         "options have failed. Both K. pneumoniae and A. baumannii appear on the WHO "
         "Global Priority Pathogen List as 'critical' threats - the highest tier. "
         "For these organisms, carbapenem resistance means virtually no remaining "
         "oral or intravenous treatment options.")
    body(doc,
         "The central failure of current surveillance is that it is binary: each "
         "isolate is classified Susceptible (S), Intermediate (I), or Resistant (R) "
         "against a fixed breakpoint. An isolate with MIC = 1 mg/L and one with "
         "MIC = 7 mg/L are both reported as 'S' against a breakpoint of 8 mg/L - "
         "but they are biologically very different. MIC creep is the gradual drift "
         "of that distribution upward, year over year. By the time enough isolates "
         "cross the breakpoint to trigger a formal resistance alert, therapeutic "
         "efficacy has already eroded for the preceding cohort.")
    body(doc,
         "Machine learning on log2(MIC) as a continuous target enables early detection "
         "of this drift before it becomes a clinical crisis - and provides quantitative "
         "forecasting that binary S/I/R reporting cannot.")
    body(doc,
         "Observed trends in our dataset: K. pneumoniae resistance (MIC > 8 mg/L) "
         "rose from 5% (2007) to 20% (2022). A. baumannii was already at 39% resistance "
         "in 2006 and reached 69% by 2022 - with MIC90 pegged at the assay panel ceiling "
         "since 2005, indicating widespread pan-resistance.")

    # ── 3. Data ───────────────────────────────────────────────────────────────
    heading(doc, "3. Data Sources and Description")
    body(doc,
         "Source: ATLAS (Antimicrobial Testing Leadership and Surveillance, Pfizer), "
         "accessed via the Vivli AMR Register under a Data Use Agreement. "
         "The dataset spans 2004-2022 across 100+ countries. Raw data are stored locally "
         "and never committed to version control per Vivli data stewardship terms.")

    body(doc, "Table 1: Dataset summary")
    add_table(doc,
        headers=["Parameter", "K. pneumoniae", "A. baumannii"],
        rows=[
            ["Total isolates", "89,572", "37,540"],
            ["Countries", "81", "79"],
            ["Train period", "2004-2018 (n=61,780)", "2004-2018 (n=23,990)"],
            ["Test period", "2019-2022 (n=27,792)", "2019-2022 (n=13,550)"],
            ["% censored MIC", "~75%", "~10%"],
            ["% resistant (MIC>8 mg/L)", "~7.4%", "~24.7%"],
            ["MIC90 trend", "+1.97 mg/L/yr (R2=0.67)", "At ceiling since 2005"],
        ],
        col_widths=[2.2, 2.0, 2.0]
    )

    heading(doc, "3.1 Variables Used", level=2)
    body(doc,
         "Raw ATLAS columns used and their derived features:")
    body(doc, "Table 2: Feature variables")
    add_table(doc,
        headers=["Feature", "Type", "Derived from", "Notes"],
        rows=[
            ["year", "Continuous", "Collection year", "Primary MIC creep driver"],
            ["gender_male", "Binary (0/1)", "Gender field", "Female is reference"],
            ["age_paediatric", "Binary (0/1)", "Age Group 0-17", "Adult (18-60) is reference"],
            ["age_elderly", "Binary (0/1)", "Age Group 61+", ""],
            ["military_proxy", "Binary (0/1)", "wound + male + 18-60", "Combat infection proxy"],
            ["spec_wound / blood / respiratory / urine / peritoneal", "Binary OHE (5)", "Source field", "Other is reference"],
            ["ctry_*", "Binary OHE (81/79)", "Country field", "Top-N countries + Other"],
            ["KPC_pos, NDM_pos, OXA_pos, VIM_pos, IMP_pos, GES_pos", "Binary (0/1)", "PCR resistance genes", "ATLAS panel"],
            ["is_censored", "Binary (0/1)", "MIC operator (>, <=)", "Data artifact flag"],
            ["pct_censored_year", "Float [0,1]", "Computed per year", "Controls panel methodology shifts"],
        ],
        col_widths=[1.6, 1.0, 1.4, 1.6]
    )
    body(doc, "Target variable: log2(MIC_numeric) - continuous regression target.")

    heading(doc, "3.2 Data Cleaning Procedures", level=2)
    body(doc,
         "MIC string parsing: Raw ATLAS MIC values are stored as strings with comparison "
         "operators (e.g., '>8', '<=0.5', '0.25'). A regex parser extracts the operator "
         "and numeric value. Unparseable entries (malformed strings, nulls) are dropped.")
    body(doc,
         "Censored MIC imputation: When the assay panel cannot bound the true MIC, labs "
         "report censored values. These are resolved using boundary doubling-dilution logic: "
         ">8 -> 16 (one step above ceiling), <=0.5 -> 0.25 (one step below floor). "
         "This is a conservative approximation - it systematically underestimates true MIC "
         "for highly resistant isolates, which is acknowledged as a limitation.")
    body(doc,
         "Log2 transformation: MIC values follow a log-normal distribution due to the "
         "doubling-dilution assay design. Log2 transformation produces a near-symmetric "
         "distribution suitable for regression. Each integer unit in log2 space corresponds "
         "to one doubling dilution step (a clinically standard unit of comparison).")
    body(doc,
         "Specimen type standardization: The free-text Source field is mapped to 5 "
         "categories (wound, blood, respiratory, urine, peritoneal) using substring matching. "
         "Unmatched entries are labeled 'other' (reference category).")
    body(doc,
         "Country encoding: All countries appearing in the training set are one-hot encoded. "
         "Countries appearing only in the test set (or infrequently) are grouped as the "
         "reference category via drop_first encoding. This handles novel surveillance regions "
         "without breaking the model.")
    body(doc,
         "Dataset filter: Only Meropenem measurements are retained. Merck/SMART data "
         "were excluded per Vivli challenge rules.")

    heading(doc, "3.3 Missing Value Treatment", level=2)
    body(doc,
         "MIC value: rows with null or unparseable MIC are dropped (~0.3% of records). "
         "No imputation - a null MIC cannot be meaningfully approximated.")
    body(doc,
         "Gender: missing gender is treated as a separate category (not imputed). "
         "The gender_male binary flag is 0 for both female and unknown.")
    body(doc,
         "Age group: missing age is treated as 'adult' (reference category) - "
         "both age_paediatric and age_elderly flags are 0.")
    body(doc,
         "Country: missing or unrecognized country is absorbed into the reference "
         "category (no separate OHE column created). The model does not see it as "
         "a distinct signal.")
    body(doc,
         "Specimen source: NaN source is mapped to 'other' (reference category) "
         "by the map_specimen function.")
    body(doc,
         "Resistance genes: missing gene PCR result in ATLAS is treated as 0 (not detected). "
         "This is a known limitation - absence of a PCR result and true gene absence are "
         "treated identically.")

    heading(doc, "3.4 Descriptive Statistics", level=2)
    body(doc,
         "The MIC distributions for the two species are structurally different, which "
         "drives diverging model performance. K. pneumoniae has a strongly bimodal "
         "distribution: ~75% of isolates pile at the assay floor (censored, MIC <= 0.25 "
         "mg/L) because most strains are fully susceptible. A. baumannii has a much more "
         "spread distribution with ~25% resistance prevalence, closer to a unimodal log-normal.")
    body(doc, "Table 3: Key descriptive statistics (log2 MIC, test set 2019-2022)")
    add_table(doc,
        headers=["Statistic", "K. pneumoniae", "A. baumannii"],
        rows=[
            ["Test set size", "27,792", "13,550"],
            ["% resistant (MIC>8 mg/L)", "~16.5%", "~68.5%"],
            ["Resistant isolates (n)", "4,588", "9,272"],
            ["% censored MIC", "~75%", "~10%"],
            ["log2(MIC) range", "-2 to +5", "-2 to +5"],
            ["Dominant mode", "log2(0.25) = -2", "log2(8) = +3"],
        ],
        col_widths=[2.2, 2.0, 2.0]
    )
    body(doc,
         "EDA charts (MIC distribution histograms, year-over-year MIC90 trends, "
         "resistance prevalence by country) are available in reports/eda/ and "
         "reports/model/.")

    # ── 4. Methodology ────────────────────────────────────────────────────────
    heading(doc, "4. Methodology")

    heading(doc, "4.1 Feature Engineering", level=2)
    body(doc,
         "Features: year (ordinal), country (one-hot encoded, top-N + 'Other'), "
         "age_group (paediatric/adult/elderly), sex (binary), specimen_type (wound/blood/"
         "urine/respiratory/other), is_censored (data artifact flag), "
         "pct_censored_year (rolling censoring rate per year). "
         "The target variable is log2(MIC_numeric).")

    heading(doc, "4.2 Time-Aware Train/Test Split", level=2)
    body(doc,
         "Data is split temporally: 2004-2018 for training, 2019-2022 for testing. "
         "Random shuffling is explicitly disabled to prevent future data leakage. "
         "This split approximates a real-world deployment scenario where the model "
         "is trained on historical data and evaluated on unseen future isolates.")

    heading(doc, "4.3 Models", level=2)
    body(doc,
         "Three models are trained per species, representing a spectrum from interpretable to accurate. "
         "Linear Regression (baseline) provides directly readable year coefficients as the MIC creep rate. "
         "Random Forest (200 trees, sklearn defaults) serves as a robust non-linear baseline. "
         "XGBoost Regressor with Optuna hyperparameter tuning (60 trials, 5-fold time-aware CV) "
         "is the primary model. All three are evaluated on: (a) all test isolates, "
         "and (b) the clinically critical resistant subset (MIC > 8 mg/L, EUCAST breakpoint).")

    heading(doc, "4.4 Explainability", level=2)
    body(doc,
         "SHAP (SHapley Additive exPlanations) values are computed for both models "
         "to explain feature contributions to individual predictions. "
         "Beeswarm plots visualize the direction and magnitude of each feature's "
         "influence. Outputs were validated for biological plausibility by the domain expert.")

    # ── 5. Results ────────────────────────────────────────────────────────────
    heading(doc, "5. Results")

    body(doc, "Table 2: Model Performance on Test Set (2019-2022)")
    add_table(doc,
        headers=["Species", "Model", "RMSE (all)", "MAE (all)", "R2 (all)", "RMSE (R)", "MAE (R)"],
        rows=[
            ["K. pneumoniae", "Linear Regression", "~2.10",  "~1.65",  "~0.05", "~4.18",  "~3.40"],
            ["K. pneumoniae", "Random Forest",     "1.558",  "1.109",  "0.21",  "2.869",  "2.180"],
            ["K. pneumoniae", "XGBoost Tuned",     "1.758",  "1.002",  "-0.01", "1.960",  "1.127"],
            ["A. baumannii",  "Linear Regression", "~1.72",  "~1.20",  "~0.08", "~1.38",  "~0.92"],
            ["A. baumannii",  "Random Forest",     "1.338",  "0.789",  "0.46",  "0.983",  "0.510"],
            ["A. baumannii",  "XGBoost Tuned",     "1.379",  "0.707",  "0.43",  "0.748",  "0.270"],
        ],
        col_widths=[1.4, 1.5, 0.85, 0.85, 0.65, 0.85, 0.85]
    )

    body(doc,
         "RMSE (R) = RMSE on resistant isolates only (MIC > 8 mg/L). "
         "LR metrics are estimated from pre-computed fallback values pending full pipeline run. "
         "For K. pneumoniae, the bimodal MIC distribution (~75% at censoring floor) "
         "causes R2 near zero on the full set; the resistant-subset RMSE is the "
         "clinically meaningful metric. "
         "XGBoost significantly outperforms both LR and RF on resistant isolates: "
         "K. pneumoniae RMSE 1.960 vs 2.869 vs ~4.18; A. baumannii RMSE 0.748 vs 0.983 vs ~1.38.")

    body(doc,
         "Top SHAP features (both species): year, pct_censored_year, is_censored, "
         "country, specimen_type. The year feature shows a consistent positive "
         "SHAP contribution in post-2015 isolates, confirming a quantifiable upward "
         "drift in MIC values independent of geographic and demographic factors.")

    # ── 6. Discussion ─────────────────────────────────────────────────────────
    heading(doc, "6. Discussion")
    body(doc,
         "For K. pneumoniae, the extremely high proportion of censored MIC values (~75%) "
         "produces a bimodal distribution that limits overall R2. This is a data artifact, "
         "not a modeling failure - most isolates simply cannot be resolved to a precise "
         "MIC by the assay panel. The resistant subset (7.4% of isolates) shows "
         "meaningful predictive signal (RMSE 1.96 log2-units with XGBoost).")

    body(doc,
         "For A. baumannii, the substantially higher resistance prevalence (24.7%) and "
         "lower censoring rate produce much cleaner signal. XGBoost achieves RMSE 0.748 "
         "on resistant isolates, approaching one doubling dilution precision - clinically "
         "actionable accuracy.")

    body(doc,
         "Limitations: (1) ATLAS reflects hospital settings in participating countries, "
         "introducing selection bias. (2) Censored MIC imputation introduces systematic "
         "downward bias at distribution tails. (3) Country is encoded from historical "
         "data and may not generalize to novel surveillance regions. "
         "(4) Year is the strongest predictor in SHAP, which may partially reflect "
         "changing assay panel ceilings over time rather than true biological resistance drift.")

    # ── 7. Conclusions ────────────────────────────────────────────────────────
    heading(doc, "7. Conclusions")
    body(doc,
         "This project demonstrates that machine learning can detect and quantify MIC "
         "Creep in large-scale AMR surveillance datasets. XGBoost with time-aware "
         "validation achieves clinically meaningful accuracy on resistant isolates, "
         "particularly for A. baumannii (RMSE 0.748 log2-units). The deployed pipeline "
         "- from raw ATLAS CSV to JSON API artefacts - runs end-to-end via a single "
         "bash script (run_all.sh). The interactive dashboard (Vercel + FastAPI on Render) "
         "provides a public-facing interface for trend visualization without exposing "
         "any raw isolate-level data.")

    body(doc,
         "This PoC establishes the feasibility of automated MIC Creep surveillance "
         "and provides a reproducible foundation for future extension to all WHO priority "
         "pathogens.")

    # ── 8. Reproducibility ────────────────────────────────────────────────────
    heading(doc, "8. Reproducibility and Source Code")
    body(doc,
         "All code is open-source: https://github.com/AnnGur/mic-creep-predict\n"
         "Pipeline: bash run_all.sh (requires ATLAS raw data in data/raw/)\n"
         "Report notebook: notebooks/final_report.ipynb (Colab-compatible)\n"
         "Model artefacts: Hugging Face Hub (loaded by FastAPI at startup)\n"
         "Raw data: cannot be redistributed per Vivli DUA; access via Vivli AMR Register.")

    # ── 9. Contribution Statement ─────────────────────────────────────────────
    heading(doc, "9. Contribution Statement")
    body(doc,
         "Anna Gurina (Technical Lead): data pipeline, feature engineering, model training "
         "and tuning, SHAP analysis, FastAPI backend, Next.js frontend, infrastructure "
         "deployment, notebook report.\n"
         "Inna Kucherova (Research Lead / Domain Expert): biological interpretation of "
         "MIC distributions and SHAP outputs, group definitions (military/paediatric proxies), "
         "clinical plausibility review, submission narrative.")

    # ── 10. References ────────────────────────────────────────────────────────
    heading(doc, "10. References")
    refs = [
        "[1] Pfizer ATLAS. Antimicrobial Testing Leadership and Surveillance. "
        "Vivli AMR Register, 2004-2022.",
        "[2] EUCAST. Clinical Breakpoints - Bacteria v14.0. "
        "European Committee on Antimicrobial Susceptibility Testing, 2024.",
        "[3] Chen T, Guestrin C. XGBoost: A Scalable Tree Boosting System. "
        "KDD 2016.",
        "[4] Lundberg SM, Lee S-I. A Unified Approach to Interpreting Model Predictions. "
        "NeurIPS 2017.",
        "[5] Akiba T, et al. Optuna: A Next-generation Hyperparameter Optimization "
        "Framework. KDD 2019.",
    ]
    for ref in refs:
        p = doc.add_paragraph(ref, style="List Bullet")
        for run in p.runs:
            set_font(run, size=10)

    # ── Save ──────────────────────────────────────────────────────────────────
    doc.save(OUT)
    print(f"Saved: {OUT}")


if __name__ == "__main__":
    build_report()
